import re
import logging
from pathlib import Path
from typing import Any

from docx import Document


logger = logging.getLogger(__name__)

CHAPTER_PATTERN = re.compile(r"\b(?:BAB|CHAPTER)\s+([1-5IVX]+)\b", re.IGNORECASE)
CITATION_PATTERN = re.compile(
    r"\(([A-Z][A-Za-zÀ-ÿ'`\-]+(?:\s+et\s+al\.)?(?:\s*&\s*[A-Z][A-Za-zÀ-ÿ'`\-]+)?),\s*(\d{4}[a-z]?)\)"
)
NARRATIVE_CITATION_PATTERN = re.compile(
    r"\b([A-Z][A-Za-zÀ-ÿ'`\-]+(?:\s+et\s+al\.)?(?:\s+and\s+[A-Z][A-Za-zÀ-ÿ'`\-]+)?)\s*\((\d{4}[a-z]?)\)"
)
BRACKET_CITATION_PATTERN = re.compile(r"\(([^()]*\b(?:19|20)\d{2}[a-z]?[^()]*)\)")
TABLE_CAPTION_PATTERN = re.compile(r"^\s*(?:Jadual|Table)\s+\d+(?:\.\d+)*\b.*", re.IGNORECASE)
OBJECTIVE_PATTERN = re.compile(
    r"\b(?:objektif|objective|research objective|RO)\s*(?:kajian|research)?\s*\d*[:.)-]?\s*(.+)",
    re.IGNORECASE,
)
OBJECTIVE_SECTION_PATTERN = re.compile(
    r"\b(?:objektif\s+kajian|objektif\s+penyelidikan|research\s+objectives?)\b",
    re.IGNORECASE,
)
FALSE_OBJECTIVE_PATTERN = re.compile(
    r"\b(?:learning\s+objectives?|module\s+objectives?|training\s+objectives?|assessment\s+objectives?|objektif\s+pembelajaran|objektif\s+modul|objektif\s+latihan|objektif\s+penilaian)\b",
    re.IGNORECASE,
)
ITEM_PREFIX_PATTERN = re.compile(r"^\s*(?:\d+[.)]|[ivxlcdm]+[.)]|[a-z][.)]|[-•])\s+(.+)", re.IGNORECASE)
NUMBERED_OBJECTIVE_PATTERN = re.compile(r"^\s*(?:\d+[.)]|[ivxlcdm]+[.)])\s+(.+)", re.IGNORECASE)
BULLET_OBJECTIVE_PATTERN = re.compile(r"^\s*(?:[-•])\s+(.+)", re.IGNORECASE)
PHASE_OBJECTIVE_PATTERN = re.compile(r"^\s*(Fasa\s+[1-3]|Phase\s+[1-3])\s*[:.-]\s*(.+)", re.IGNORECASE)
OBJECTIVE_SPECIFIC_SUBHEADING_PATTERN = re.compile(r"^\s*(?:objektif\s+khusus|specific\s+objectives?)\s*$", re.IGNORECASE)
OBJECTIVE_GENERAL_SUBHEADING_PATTERN = re.compile(r"^\s*(?:objektif\s+umum|general\s+objectives?)\s*$", re.IGNORECASE)
OBJECTIVE_STOP_HEADING_PATTERN = re.compile(
    r"\b(?:persoalan\s+kajian|research\s+questions?|hipotesis|hypothes(?:is|es)|kepentingan\s+kajian|skop\s+kajian|sorotan\s+literatur|literature\s+review)\b",
    re.IGNORECASE,
)
REFERENCE_HEADING_PATTERN = re.compile(r"^\s*(?:references|rujukan|bibliografi|bibliography)\s*$", re.IGNORECASE)
HEADING_PATTERNS = [
    re.compile(r"^\s*(?:BAB|CHAPTER)\s+[1-5IVX]+\b.*", re.IGNORECASE),
    re.compile(r"^\s*\d+(?:\.\d+){0,4}\s+[A-Z0-9].{2,120}$"),
    re.compile(
        r"^\s*(?:PENGENALAN|INTRODUCTION|LITERATURE REVIEW|SOROTAN LITERATUR|METHODOLOGY|METODOLOGI|DAPATAN|FINDINGS|PERBINCANGAN|DISCUSSION|KESIMPULAN|CONCLUSION)\s*$",
        re.IGNORECASE,
    ),
]


def parse_uploaded_documents(upload_dir: Path) -> dict[str, Any]:
    parsed_documents: list[dict[str, Any]] = []
    aggregate = {
        "chapters": [],
        "headings": [],
        "paragraphs": [],
        "tables": [],
        "citations": [],
        "references": [],
        "objectives": [],
        "objective_candidates": [],
        "table_captions": [],
    }

    for file_path in sorted(upload_dir.iterdir()):
        if not file_path.is_file() or file_path.name.startswith("."):
            continue

        suffix = file_path.suffix.lower()
        if suffix == ".docx":
            document_result = parse_docx(file_path)
        elif suffix == ".pdf":
            document_result = parse_pdf(file_path)
        else:
            continue

        parsed_documents.append(document_result)
        for key in aggregate:
            aggregate[key].extend(document_result.get(key, []))

    normalized = {key: deduplicate_items(value) for key, value in aggregate.items()}
    objective_result = extract_research_objectives_result(
        parsed_documents,
        normalized["objective_candidates"],
    )
    normalized["objectives"] = objective_result["objectives"]

    return {
        "project_files_parsed": len(parsed_documents),
        "documents": parsed_documents,
        "objective_extraction_status": objective_result["objective_extraction_status"],
        "general_objective": objective_result["general_objective"],
        "rejected_objective_candidates": objective_result["rejected_objective_candidates"],
        "objective_extraction_metadata": objective_result["objective_extraction_metadata"],
        "objective_debug": objective_result.get("objective_debug", {}),
        **normalized,
    }


def parse_docx(file_path: Path) -> dict[str, Any]:
    document = Document(file_path)
    headings: list[dict[str, Any]] = []
    paragraphs: list[dict[str, Any]] = []
    chapters: list[dict[str, Any]] = []
    citations: list[dict[str, Any]] = []
    references: list[dict[str, Any]] = []
    objective_candidates: list[dict[str, Any]] = []
    table_captions: list[dict[str, Any]] = []

    in_references = False

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = normalize_text(paragraph.text)
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else ""
        is_heading = style_name.lower().startswith("heading")
        if is_heading:
            headings.append(
                {
                    "source_file": file_path.name,
                    "text": text,
                    "style": style_name,
                    "position": index,
                }
            )

        chapter_match = CHAPTER_PATTERN.search(text)
        if chapter_match:
            chapters.append(
                {
                    "source_file": file_path.name,
                    "label": chapter_match.group(0),
                    "text": text,
                    "position": index,
                }
            )

        if REFERENCE_HEADING_PATTERN.match(text):
            in_references = True
            continue

        if in_references:
            references.append(
                {
                    "source_file": file_path.name,
                    "text": text,
                    "position": index,
                }
            )
            continue

        if TABLE_CAPTION_PATTERN.match(text):
            table_captions.append(
                {
                    "source_file": file_path.name,
                    "caption": text,
                    "position": index,
                }
            )

        objective_match = OBJECTIVE_PATTERN.search(text)
        if objective_match:
            objective_candidates.append(
                {
                    "source_file": file_path.name,
                    "text": text,
                    "detected_objective": objective_match.group(1).strip(),
                    "position": index,
                }
            )

        paragraphs.append(
            {
                "source_file": file_path.name,
                "text": text,
                "style": style_name,
                "position": index,
            }
        )
        citations.extend(extract_citations(text, file_path.name, index))

    tables = extract_docx_tables(document, file_path.name)

    return {
        "source_file": file_path.name,
        "file_type": "docx",
        "chapters": chapters,
        "headings": headings,
        "paragraphs": paragraphs,
        "tables": tables,
        "citations": deduplicate_items(citations),
        "references": references,
        "objectives": [],
        "objective_candidates": objective_candidates,
        "table_captions": table_captions,
    }


def parse_pdf(file_path: Path) -> dict[str, Any]:
    page_texts: list[dict[str, Any]] = []
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(file_path))
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                page_texts.append({"page": page_number, "text": text})
    except Exception as error:
        return empty_pdf_result(file_path, parser_status="failed", message=str(error))

    if not page_texts:
        return empty_pdf_result(
            file_path,
            parser_status="no_text_detected",
            message="PDF was read, but no extractable text was detected.",
        )

    chapters: list[dict[str, Any]] = detect_chapters(file_path.name, page_texts)
    headings: list[dict[str, Any]] = []
    paragraphs: list[dict[str, Any]] = []
    citations: list[dict[str, Any]] = []
    references: list[dict[str, Any]] = []
    objective_candidates: list[dict[str, Any]] = []
    table_captions: list[dict[str, Any]] = []

    in_references = False
    position = 0

    for page_text in page_texts:
        page_number = int(page_text["page"])
        for line in split_pdf_lines(str(page_text["text"])):
            text = normalize_text(line)
            if not text:
                continue

            position += 1

            if is_likely_heading(text):
                headings.append(
                    {
                        "source_file": file_path.name,
                        "text": text,
                        "style": "pdf_detected_heading",
                        "page": page_number,
                        "position": position,
                    }
                )

            chapter_match = CHAPTER_PATTERN.search(text)
            if chapter_match:
                chapters.append(
                    {
                        "source_file": file_path.name,
                        "label": chapter_match.group(0),
                        "text": text,
                        "page": page_number,
                        "position": position,
                    }
                )

            if REFERENCE_HEADING_PATTERN.match(text):
                in_references = True
                continue

            if in_references:
                references.append(
                    {
                        "source_file": file_path.name,
                        "text": text,
                        "page": page_number,
                        "position": position,
                    }
                )
                continue

            if TABLE_CAPTION_PATTERN.match(text):
                table_captions.append(
                    {
                        "source_file": file_path.name,
                        "caption": text,
                        "page": page_number,
                        "position": position,
                    }
                )

            objective_match = OBJECTIVE_PATTERN.search(text)
            if objective_match:
                objective_candidates.append(
                    {
                        "source_file": file_path.name,
                        "text": text,
                        "detected_objective": objective_match.group(1).strip(),
                        "page": page_number,
                        "position": position,
                    }
                )

            paragraphs.append(
                {
                    "source_file": file_path.name,
                    "text": text,
                    "style": "pdf_text",
                    "page": page_number,
                    "position": position,
                }
            )
            citations.extend(extract_citations(text, file_path.name, position))

    return {
        "source_file": file_path.name,
        "file_type": "pdf",
        "parser_status": "parsed",
        "message": "PDF text extracted using PyPDF2.",
        "page_count": len(page_texts),
        "chapters": deduplicate_items(chapters),
        "headings": deduplicate_items(headings),
        "paragraphs": paragraphs,
        "tables": [],
        "citations": deduplicate_items(citations),
        "references": references,
        "objectives": [],
        "objective_candidates": deduplicate_items(objective_candidates),
        "table_captions": table_captions,
    }


def empty_pdf_result(file_path: Path, parser_status: str, message: str) -> dict[str, Any]:
    return {
        "source_file": file_path.name,
        "file_type": "pdf",
        "parser_status": parser_status,
        "message": message,
        "page_count": 0,
        "chapters": detect_chapters(file_path.name, []),
        "headings": [],
        "paragraphs": [],
        "tables": [],
        "citations": [],
        "references": [],
        "objectives": [],
        "objective_candidates": [],
        "table_captions": [],
    }


def detect_chapters(source_file: str, page_texts: list[dict[str, Any]]) -> list[dict[str, Any]]:
    chapters: list[dict[str, Any]] = []
    filename_text = source_file.replace("_", " ").replace("-", " ")
    filename_match = CHAPTER_PATTERN.search(filename_text)
    if filename_match:
        chapters.append(
            {
                "source_file": source_file,
                "label": filename_match.group(0),
                "text": filename_text,
                "position": 0,
            }
        )

    for page_text in page_texts:
        match = CHAPTER_PATTERN.search(str(page_text.get("text", ""))[:1000])
        if match:
            chapters.append(
                {
                    "source_file": source_file,
                    "label": match.group(0),
                    "text": match.group(0),
                    "page": page_text.get("page", 0),
                    "position": 0,
                }
            )

    return deduplicate_items(chapters)


def split_pdf_lines(text: str) -> list[str]:
    rough_lines = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9(])|\n+", text)
    return [line.strip() for line in rough_lines if line.strip()]


def is_likely_heading(text: str) -> bool:
    if len(text) > 140:
        return False

    if any(pattern.match(text) for pattern in HEADING_PATTERNS):
        return True

    words = text.split()
    if 2 <= len(words) <= 12 and text.isupper():
        return True

    return False


def extract_research_objectives(
    documents: list[dict[str, Any]],
    objective_candidates: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], str]:
    chapter_one_documents = [document for document in documents if is_chapter_one_document(document)]
    search_documents = chapter_one_documents or documents
    extracted: list[dict[str, Any]] = []

    for document in search_documents:
        source_file = str(document.get("source_file", "Unknown file"))
        paragraphs = [item for item in document.get("paragraphs", []) if isinstance(item, dict)]
        headings = [item for item in document.get("headings", []) if isinstance(item, dict)]
        objective_heading = find_objective_heading(headings, paragraphs)
        if not objective_heading:
            continue

        items = extract_objective_items_below_heading(source_file, objective_heading, paragraphs, headings)
        logger.info(
            "objective section found: source_file=%s selected_heading=%s objective_count=%s objective_source_page=%s",
            source_file,
            objective_heading.get("text", ""),
            len(items),
            objective_heading.get("page", ""),
        )
        extracted.extend(items)

    filtered = deduplicate_items([objective for objective in extracted if is_valid_research_objective(objective.get("objective_text", ""))])
    if filtered:
        logger.info("objective items extracted: objective_count=%s", len(filtered[:10]))
        return assign_objective_ids(filtered[:10], extraction_status="extracted"), "extracted"

    fallback = build_fallback_objectives(search_documents, objective_candidates)
    logger.info("objective extraction fallback used: count=%s", len(fallback))
    return fallback, "fallback"


def is_chapter_one_document(document: dict[str, Any]) -> bool:
    source_file = str(document.get("source_file", "")).lower().replace("_", " ")
    if re.search(r"\b(?:bab|chapter)\s*1\b", source_file):
        return True

    chapters = document.get("chapters", [])
    if isinstance(chapters, list):
        for chapter in chapters:
            if isinstance(chapter, dict) and re.search(r"\b(?:bab|chapter)\s*1\b", str(chapter.get("label", "")).lower()):
                return True

    return False


def find_objective_heading(
    headings: list[dict[str, Any]],
    paragraphs: list[dict[str, Any]],
) -> dict[str, Any] | None:
    for heading in headings:
        text = str(heading.get("text", ""))
        if is_objective_section_text(text):
            return heading

    for paragraph in paragraphs:
        text = str(paragraph.get("text", ""))
        if is_objective_section_text(text) and len(text) <= 180:
            return paragraph

    return None


def is_objective_section_text(text: str) -> bool:
    return bool(OBJECTIVE_SECTION_PATTERN.search(text)) and not FALSE_OBJECTIVE_PATTERN.search(text)


def extract_objective_items_below_heading(
    source_file: str,
    objective_heading: dict[str, Any],
    paragraphs: list[dict[str, Any]],
    headings: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    heading_position = int(objective_heading.get("position", 0) or 0)
    next_heading_position = find_next_major_heading_position(source_file, heading_position, headings)
    source_heading = str(objective_heading.get("text", "Research Objectives"))
    section_paragraphs = collect_objective_section_paragraphs(
        source_file,
        heading_position,
        next_heading_position,
        paragraphs,
    )
    selected_subheading = select_objective_subheading(section_paragraphs)
    if selected_subheading:
        logger.info(
            "selected_subheading=%s source_file=%s objective_source_page=%s",
            selected_subheading.get("text", ""),
            source_file,
            selected_subheading.get("page", ""),
        )
        section_paragraphs = [
            paragraph
            for paragraph in section_paragraphs
            if int(paragraph.get("position", 0) or 0) > int(selected_subheading.get("position", 0) or 0)
        ]
        source_heading = str(selected_subheading.get("text", source_heading))

    phase_items = extract_phase_objectives(source_file, source_heading, section_paragraphs)
    if phase_items:
        return phase_items

    numbered_items = extract_prefixed_objectives(
        source_file,
        source_heading,
        section_paragraphs,
        NUMBERED_OBJECTIVE_PATTERN,
        confidence_score=95,
    )
    if numbered_items:
        return numbered_items

    return extract_prefixed_objectives(
        source_file,
        source_heading,
        section_paragraphs,
        BULLET_OBJECTIVE_PATTERN,
        confidence_score=80,
    )


def collect_objective_section_paragraphs(
    source_file: str,
    heading_position: int,
    next_heading_position: int | None,
    paragraphs: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    collected = []
    for paragraph in paragraphs:
        if str(paragraph.get("source_file", source_file)) != source_file:
            continue
        position = int(paragraph.get("position", 0) or 0)
        if position <= heading_position:
            continue
        if next_heading_position and position >= next_heading_position:
            break
        text = str(paragraph.get("text", ""))
        if OBJECTIVE_STOP_HEADING_PATTERN.search(text):
            break
        collected.append(paragraph)
    return collected


def select_objective_subheading(section_paragraphs: list[dict[str, Any]]) -> dict[str, Any] | None:
    for paragraph in section_paragraphs:
        if OBJECTIVE_SPECIFIC_SUBHEADING_PATTERN.match(str(paragraph.get("text", ""))):
            return paragraph
    return None


def extract_phase_objectives(
    source_file: str,
    source_heading: str,
    section_paragraphs: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    items = []
    for paragraph in section_paragraphs:
        text = str(paragraph.get("text", ""))
        match = PHASE_OBJECTIVE_PATTERN.match(text)
        if not match:
            continue
        item_text = clean_objective_item(match.group(2), allow_inferred=True)
        if item_text:
            items.append(build_objective_item(source_file, source_heading, paragraph, item_text, 90))
    return items


def extract_prefixed_objectives(
    source_file: str,
    source_heading: str,
    section_paragraphs: list[dict[str, Any]],
    pattern: re.Pattern[str],
    confidence_score: int,
) -> list[dict[str, Any]]:
    items = []
    ignore_until_specific = any(
        OBJECTIVE_SPECIFIC_SUBHEADING_PATTERN.match(str(paragraph.get("text", "")))
        for paragraph in section_paragraphs
    )
    specific_started = not ignore_until_specific

    for paragraph in section_paragraphs:
        text = str(paragraph.get("text", ""))
        if OBJECTIVE_SPECIFIC_SUBHEADING_PATTERN.match(text):
            specific_started = True
            continue
        if OBJECTIVE_GENERAL_SUBHEADING_PATTERN.match(text):
            specific_started = False if ignore_until_specific else specific_started
            continue
        if not specific_started or FALSE_OBJECTIVE_PATTERN.search(text) or looks_like_citation_fragment(text):
            continue

        match = pattern.match(text)
        if not match:
            continue
        item_text = clean_objective_item(match.group(1), allow_inferred=False)
        if item_text:
            items.append(build_objective_item(source_file, source_heading, paragraph, item_text, confidence_score))
    return items


def build_objective_item(
    source_file: str,
    source_heading: str,
    paragraph: dict[str, Any],
    item_text: str,
    confidence_score: int,
) -> dict[str, Any]:
    return {
        "objective_text": item_text,
        "source_file": source_file,
        "source_chapter": "Bab 1",
        "source_heading": source_heading,
        "position": int(paragraph.get("position", 0) or 0),
        "page": paragraph.get("page", ""),
        "confidence_score": confidence_score,
        "extraction_status": "extracted",
        "selected_heading": source_heading,
        "selected_subheading": source_heading,
        "objective_source_page": paragraph.get("page", ""),
    }


def find_next_major_heading_position(
    source_file: str,
    heading_position: int,
    headings: list[dict[str, Any]],
) -> int | None:
    later_headings = []
    for heading in headings:
        if str(heading.get("source_file", source_file)) != source_file:
            continue
        position = int(heading.get("position", 0) or 0)
        text = str(heading.get("text", ""))
        if position > heading_position and OBJECTIVE_STOP_HEADING_PATTERN.search(text):
            later_headings.append(position)
        elif position > heading_position and is_likely_major_heading(text):
            later_headings.append(position)

    return min(later_headings) if later_headings else None


def is_likely_major_heading(text: str) -> bool:
    return bool(HEADING_PATTERNS[0].match(text) or re.match(r"^\s*\d+(?:\.\d+){0,3}\s+\S+", text))


def clean_objective_item(text: str, allow_inferred: bool = False) -> str:
    stripped = normalize_text(text)
    match = ITEM_PREFIX_PATTERN.match(stripped)
    if match:
        stripped = match.group(1).strip()
    stripped = re.sub(r"^(?:untuk|to)\s+", "", stripped, flags=re.IGNORECASE)
    if len(stripped) < 18 or len(stripped) > 320:
        return ""
    if not allow_inferred and not re.search(r"\b(?:mengenal pasti|mengenalpasti|menentukan|menganalisis|membangunkan|menilai|mengkaji|to identify|to determine|to analyse|to analyze|to develop|to evaluate|identify|determine|analyse|analyze|develop|evaluate)\b", stripped, re.IGNORECASE):
        return ""
    return stripped


def extract_inline_objective(text: str) -> str:
    match = re.search(r"(?:adalah|are|untuk|to)\s+(.+)$", text, re.IGNORECASE)
    if not match:
        return ""
    return clean_objective_item(match.group(1), allow_inferred=True)


def looks_like_citation_fragment(text: str) -> bool:
    if len(text) < 80 and YEAR_PATTERN_COUNT(text) >= 1:
        return True
    return text.count("(") > 2 or text.count(")") > 2


def YEAR_PATTERN_COUNT(text: str) -> int:
    return len(re.findall(r"\b(?:19|20)\d{2}[a-z]?\b", text))


def is_valid_research_objective(text: str) -> bool:
    if not text or FALSE_OBJECTIVE_PATTERN.search(text) or looks_like_citation_fragment(text):
        return False
    return bool(clean_objective_item(text, allow_inferred=True))


def build_fallback_objectives(
    documents: list[dict[str, Any]],
    objective_candidates: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    thesis_title = infer_thesis_title(documents)
    fallback_actions = [
        "Mengenal pasti keperluan kajian berdasarkan skop tesis",
        "Menganalisis elemen utama kajian berdasarkan dokumen tesis",
        "Membangunkan asas pemetaan objektif kepada dapatan kajian",
    ]
    if thesis_title:
        fallback_actions = [
            f"Mengenal pasti keperluan kajian berkaitan {thesis_title}",
            f"Menganalisis elemen utama kajian berkaitan {thesis_title}",
            f"Membangunkan asas pemetaan objektif bagi {thesis_title}",
        ]

    source_file = str(documents[0].get("source_file", "Unknown file")) if documents else "Unknown file"
    if objective_candidates:
        source_file = str(objective_candidates[0].get("source_file", source_file))

    return [
        {
            "objective_id": f"RO{index}",
            "objective_text": text,
            "source_file": source_file,
            "source_chapter": "Bab 1",
            "source_heading": "Fallback objective extraction",
            "confidence_score": 40,
            "extraction_status": "fallback",
        }
        for index, text in enumerate(fallback_actions, start=1)
    ]


def infer_thesis_title(documents: list[dict[str, Any]]) -> str:
    for document in documents:
        headings = document.get("headings", [])
        if isinstance(headings, list):
            for heading in headings:
                text = str(heading.get("text", "") if isinstance(heading, dict) else "")
                if 20 <= len(text) <= 180 and not CHAPTER_PATTERN.search(text):
                    return text
        paragraphs = document.get("paragraphs", [])
        if isinstance(paragraphs, list):
            for paragraph in paragraphs[:5]:
                text = str(paragraph.get("text", "") if isinstance(paragraph, dict) else "")
                if 20 <= len(text) <= 180 and not CHAPTER_PATTERN.search(text):
                    return text
    return ""


def assign_objective_ids(objectives: list[dict[str, Any]], extraction_status: str) -> list[dict[str, Any]]:
    assigned = []
    for index, objective in enumerate(objectives, start=1):
        assigned.append(
            {
                "objective_id": f"RO{index}",
                "objective_text": str(objective.get("objective_text", "")),
                "source_file": str(objective.get("source_file", "Unknown file")),
                "source_chapter": str(objective.get("source_chapter", "Bab 1")),
                "source_heading": str(objective.get("source_heading", "Research Objectives")),
                "confidence_score": int(objective.get("confidence_score", 80) or 80),
                "extraction_status": extraction_status,
                "selected_heading": str(objective.get("selected_heading", objective.get("source_heading", ""))),
                "selected_subheading": str(objective.get("selected_subheading", "")),
                "objective_source_page": objective.get("objective_source_page", objective.get("page", "")),
            }
        )
    return assigned


def extract_docx_tables(document: Document, source_file: str) -> list[dict[str, Any]]:
    tables = []
    for index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([normalize_text(cell.text) for cell in row.cells])

        tables.append(
            {
                "source_file": source_file,
                "table_index": index,
                "row_count": len(rows),
                "column_count": max((len(row) for row in rows), default=0),
                "rows": rows,
            }
        )

    return tables


def extract_citations(text: str, source_file: str, position: int) -> list[dict[str, Any]]:
    citations = []
    captured = set()
    for match in CITATION_PATTERN.finditer(text):
        citation = match.group(0)
        captured.add(citation)
        citations.append(build_citation_row(source_file, citation, match.group(1), match.group(2), position))

    for bracket_match in BRACKET_CITATION_PATTERN.finditer(text):
        bracket_text = bracket_match.group(0)
        if bracket_text in captured:
            continue
        inner_text = bracket_match.group(1)
        for part in inner_text.split(";"):
            parsed = parse_bracket_citation_part(part)
            if parsed:
                author, year = parsed
                citation = f"({part.strip()})"
                captured.add(citation)
                citations.append(build_citation_row(source_file, citation, author, year, position))

    for match in NARRATIVE_CITATION_PATTERN.finditer(text):
        citation = match.group(0)
        if citation not in captured:
            citations.append(build_citation_row(source_file, citation, match.group(1), match.group(2), position))

    return citations


def parse_bracket_citation_part(text: str) -> tuple[str, str] | None:
    match = re.search(r"\b((?:19|20)\d{2}[a-z]?)\b", text)
    if not match:
        return None
    author = text[: match.start()].strip(" ,")
    if not author:
        return None
    return author, match.group(1)


def build_citation_row(
    source_file: str,
    citation: str,
    author: str,
    year: str,
    position: int,
) -> dict[str, Any]:
    return {
        "source_file": source_file,
        "citation": citation,
        "author": author,
        "year": year,
        "position": position,
    }


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def deduplicate_items(items: list[Any]) -> list[Any]:
    seen = set()
    deduplicated = []
    for item in items:
        marker = repr(item)
        if marker in seen:
            continue

        seen.add(marker)
        deduplicated.append(item)

    return deduplicated


# PHASE 6.2 universal objective extractor.
UNIVERSAL_OBJECTIVE_HEADING_PATTERN = re.compile(
    r"\b(?:"
    r"objektif\s+kajian|objektif\s+penyelidikan(?:\s+kajian)?|"
    r"objektif\s+umum\s+dan\s+objektif\s+khusus|tujuan\s+kajian|matlamat\s+kajian|"
    r"matlamat\s+penyelidikan|tujuan\s+penyelidikan|"
    r"research\s+objectives?|objectives\s+of\s+the\s+study|study\s+objectives?|"
    r"general\s+objective\s+and\s+specific\s+objectives?|general\s+research\s+objective|"
    r"specific\s+research\s+objectives?|aim\s+and\s+objectives?|research\s+aim\s+and\s+objectives?"
    r")\b",
    re.IGNORECASE,
)
UNIVERSAL_SPECIFIC_SUBHEADING_PATTERN = re.compile(
    r"^\s*(?:objektif\s+khusus|specific\s+(?:research\s+)?objectives?)\s*$",
    re.IGNORECASE,
)
UNIVERSAL_GENERAL_SUBHEADING_PATTERN = re.compile(
    r"^\s*(?:objektif\s+umum|general\s+(?:research\s+)?objective)\s*$",
    re.IGNORECASE,
)
UNIVERSAL_FALSE_OBJECTIVE_PATTERN = re.compile(
    r"\b(?:"
    r"learning\s+objectives?|module\s+objectives?|course\s+objectives?|training\s+objectives?|"
    r"assessment\s+objectives?|program\s+objectives?|"
    r"objektif\s+pembelajaran|objektif\s+modul|objektif\s+latihan|objektif\s+kursus|"
    r"objektif\s+program|objektif\s+penilaian"
    r")\b",
    re.IGNORECASE,
)
UNIVERSAL_STOP_HEADING_PATTERN = re.compile(
    r"\b(?:"
    r"persoalan\s+kajian|soalan\s+kajian|hipotesis\s+kajian|kepentingan\s+kajian|"
    r"skop\s+kajian|batasan\s+kajian|definisi\s+operasional|sorotan\s+literatur|"
    r"kerangka\s+teori|research\s+questions?|hypotheses|significance\s+of\s+the\s+study|"
    r"scope\s+of\s+the\s+study|limitations|operational\s+definition|literature\s+review|"
    r"theoretical\s+framework|conceptual\s+framework"
    r")\b",
    re.IGNORECASE,
)
UNIVERSAL_PHASE_PATTERN = re.compile(r"^\s*((?:Fasa|Phase)\s+[1-3])\s*[:.)-]\s*(.+)", re.IGNORECASE)
UNIVERSAL_NUMBERED_PATTERN = re.compile(
    r"^\s*(?:(\d+)[.)]|\(([ivxlcdm]+)\)|([ivxlcdm]+)[.)]|([a-z])[.)])\s+(.+)",
    re.IGNORECASE,
)
UNIVERSAL_BULLET_PATTERN = re.compile(r"^\s*(?:[-*\u2022])\s+(.+)", re.IGNORECASE)
UNIVERSAL_OBJECTIVE_VERB_PATTERN = re.compile(
    r"\b(?:"
    r"mengenal\s+pasti|mengenalpasti|menganalisis|membangunkan|menilai|mengkaji|"
    r"menentukan|meneroka|menjelaskan|merumuskan|mengesahkan|"
    r"to\s+identify|to\s+determine|to\s+examine|to\s+investigate|to\s+assess|"
    r"to\s+evaluate|to\s+develop|to\s+explore|to\s+analyse|to\s+analyze|"
    r"to\s+test|to\s+validate|to\s+measure|to\s+compare|"
    r"identify|determine|examine|investigate|assess|evaluate|develop|explore|"
    r"analyse|analyze|test|validate|measure|compare"
    r")\b",
    re.IGNORECASE,
)


def extract_research_objectives_result(
    documents: list[dict[str, Any]],
    objective_candidates: list[dict[str, Any]],
) -> dict[str, Any]:
    chapter_one_documents = [document for document in documents if is_chapter_one_document(document)]
    search_documents = chapter_one_documents or documents
    extracted: list[dict[str, Any]] = []
    rejected: list[dict[str, Any]] = []
    general_objective: dict[str, Any] | None = None
    metadata: dict[str, Any] = {
        "selected_heading": "",
        "selected_subheading": "",
        "numbering_style_detected": "",
        "extraction_strategy": "not_found",
        "source_file": "",
        "objective_source_page": "",
        "confidence_score": 0,
    }
    objective_debug: dict[str, Any] = {
        "raw_objective_section": "",
        "selected_heading": "",
        "selected_subheading": "",
        "objective_extraction_status": "not_found",
        "reason_for_fallback": "",
        "detected_objectives": [],
        "fallback_reason": "",
        "candidate_headings_found": [],
    }
    candidate_headings_found = collect_objective_heading_candidates(documents)
    metadata["candidate_headings_found"] = candidate_headings_found
    objective_debug["candidate_headings_found"] = candidate_headings_found

    for document in search_documents:
        source_file = str(document.get("source_file", "Unknown file"))
        paragraphs = [item for item in document.get("paragraphs", []) if isinstance(item, dict)]
        headings = [item for item in document.get("headings", []) if isinstance(item, dict)]
        selected_heading = find_universal_objective_heading(headings, paragraphs, infer_document_chapter(document))
        if not selected_heading:
            continue

        section = collect_universal_objective_section(source_file, selected_heading, paragraphs, headings)
        if not section:
            objective_debug = build_objective_debug(
                selected_heading=selected_heading,
                selected_subheading=None,
                section=[],
                objectives=[],
                extraction_status="fallback",
                reason_for_fallback="objective heading found but objective section is empty",
                candidate_headings_found=candidate_headings_found,
            )
            continue

        general_objective = extract_general_objective(source_file, selected_heading, section)
        specific_section, selected_subheading = choose_specific_section(section)
        objectives, strategy, numbering_style, local_rejected = extract_specific_objectives(
            source_file=source_file,
            source_heading=str(selected_heading.get("text", "Research Objectives")),
            selected_subheading=selected_subheading,
            section=specific_section,
        )
        rejected.extend(local_rejected)
        objective_debug = build_objective_debug(
            selected_heading=selected_heading,
            selected_subheading=selected_subheading,
            section=section,
            objectives=objectives,
            extraction_status="extracted" if objectives else "fallback",
            reason_for_fallback="" if objectives else "objective section found but no numbered, bullet, or phase objectives were extracted",
            candidate_headings_found=candidate_headings_found,
        )

        metadata = {
            "selected_heading": str(selected_heading.get("text", "")),
            "selected_subheading": str(selected_subheading.get("text", "")) if selected_subheading else "",
            "numbering_style_detected": numbering_style,
            "extraction_strategy": strategy,
            "source_file": source_file,
            "objective_source_page": selected_heading.get("page", ""),
            "confidence_score": max((int(item.get("confidence_score", 0) or 0) for item in objectives), default=0),
            "candidate_headings_found": candidate_headings_found,
        }
        logger.info(
            "objective extraction metadata: selected_heading=%s selected_subheading=%s objective_count=%s objective_source_page=%s",
            metadata["selected_heading"],
            metadata["selected_subheading"],
            len(objectives),
            metadata["objective_source_page"],
        )
        log_objective_debug(objective_debug)
        extracted.extend(objectives)
        if extracted:
            break

    cleaned = [item for item in deduplicate_items(extracted) if item.get("objective_text")]
    if cleaned:
        assigned = assign_objective_ids(cleaned[:10], extraction_status="extracted")
        return {
            "objective_extraction_status": "extracted",
            "general_objective": general_objective or {},
            "objectives": assigned,
            "rejected_objective_candidates": rejected,
            "objective_extraction_metadata": metadata,
            "objective_debug": objective_debug,
        }

    fallback = build_fallback_objectives(search_documents, objective_candidates)
    rejected.extend(reject_objective_candidates(objective_candidates, "no structured research objective section found"))
    metadata["extraction_strategy"] = "fallback"
    metadata["confidence_score"] = 40
    reason_for_fallback = objective_debug.get("reason_for_fallback") or "no structured research objective section found"
    objective_debug["objective_extraction_status"] = "fallback"
    objective_debug["reason_for_fallback"] = reason_for_fallback
    objective_debug["fallback_reason"] = reason_for_fallback
    objective_debug["detected_objectives"] = fallback
    logger.info("objective extraction fallback used: count=%s", len(fallback))
    log_objective_debug(objective_debug)
    return {
        "objective_extraction_status": "fallback",
        "general_objective": general_objective or {},
        "objectives": fallback,
        "rejected_objective_candidates": rejected,
        "objective_extraction_metadata": metadata,
        "objective_debug": objective_debug,
    }


def find_universal_objective_heading(
    headings: list[dict[str, Any]],
    paragraphs: list[dict[str, Any]],
    chapter: str = "Unknown",
) -> dict[str, Any] | None:
    ranked: list[tuple[int, int, dict[str, Any]]] = []
    for item in [*headings, *paragraphs]:
        text = str(item.get("text", ""))
        if len(text) >= 120:
            continue
        score, rejected, _reason = score_objective_heading(text, chapter)
        if score <= 0 or rejected:
            continue
        position = int(item.get("position", 0) or 0)
        ranked.append((score, -position, item))
    if not ranked:
        return None
    ranked.sort(key=lambda candidate: (candidate[0], candidate[1]), reverse=True)
    return ranked[0][2]


def build_objective_debug(
    selected_heading: dict[str, Any],
    selected_subheading: dict[str, Any] | None,
    section: list[dict[str, Any]],
    objectives: list[dict[str, Any]],
    extraction_status: str,
    reason_for_fallback: str,
    candidate_headings_found: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    raw_objective_section = "\n".join(str(item.get("text", "")) for item in section if item.get("text"))
    return {
        "raw_objective_section": raw_objective_section,
        "selected_heading": str(selected_heading.get("text", "")),
        "selected_subheading": str(selected_subheading.get("text", "")) if selected_subheading else "",
        "objective_extraction_status": extraction_status,
        "reason_for_fallback": reason_for_fallback,
        "fallback_reason": reason_for_fallback,
        "detected_objectives": objectives,
        "objective_count": len(objectives),
        "objective_source_page": selected_heading.get("page", ""),
        "candidate_headings_found": candidate_headings_found or [],
    }


def log_objective_debug(objective_debug: dict[str, Any]) -> None:
    logger.info(
        "OBJECTIVE DEBUG\nDetected heading: %s\nDetected subheading: %s\nDetected objectives: %s\nFallback reason: %s",
        objective_debug.get("selected_heading", ""),
        objective_debug.get("selected_subheading", ""),
        len(objective_debug.get("detected_objectives", []) or []),
        objective_debug.get("reason_for_fallback", ""),
    )


def is_universal_objective_heading(text: str) -> bool:
    score, rejected, _reason = score_objective_heading(text)
    return score > 0 and not rejected


OBJECTIVE_HEADING_PRIORITIES: tuple[tuple[str, int], ...] = (
    ("specific research objectives", 100),
    ("objectives of the study", 100),
    ("objektif penyelidikan", 100),
    ("research objectives", 100),
    ("matlamat dan objektif kajian", 80),
    ("objektif kajian", 100),
    ("research aim and objectives", 90),
    ("aim and objectives", 90),
    ("tujuan kajian", 90),
    ("objektif umum dan objektif khusus", 85),
    ("general objective and specific objectives", 85),
    ("specific objectives", 85),
    ("objektif khusus", 85),
)


OBJECTIVE_CONTEXT_REJECT_TERMS = (
    "keselarasan",
    "hubungan",
    "pencapaian",
    "perbincangan",
    "penilaian",
    "alignment",
    "relationship",
    "achievement",
    "discussion",
    "evaluation",
)


OBJECTIVE_STRONG_REJECT_PATTERNS: tuple[re.Pattern[str], ...] = (
    re.compile(r"\bro\s*[0-9]+\b", re.IGNORECASE),
    re.compile(r"\bobjektif\s+kajian\s+(?:pertama|kedua|ketiga|keempat|kelima)\b", re.IGNORECASE),
    re.compile(r"\bmenyokong\s+objektif\b", re.IGNORECASE),
    re.compile(r"\bpencapaian\s+objektif\b", re.IGNORECASE),
    re.compile(r"\bperbincangan\s+objektif\b", re.IGNORECASE),
    re.compile(r"\bdapatan\s+objektif\b", re.IGNORECASE),
)


EXACT_OBJECTIVE_HEADINGS = {
    "objektif kajian",
    "objektif penyelidikan",
    "research objectives",
    "objectives of the study",
    "specific research objectives",
}


def collect_objective_heading_candidates(documents: list[dict[str, Any]]) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    for document in documents:
        chapter = infer_document_chapter(document)
        source_file = str(document.get("source_file", "Unknown file"))
        items = [
            item
            for item in [
                *[entry for entry in document.get("headings", []) if isinstance(entry, dict)],
                *[entry for entry in document.get("paragraphs", []) if isinstance(entry, dict)],
            ]
            if len(str(item.get("text", ""))) <= 240
        ]
        seen: set[tuple[str, int]] = set()
        for item in items:
            heading = str(item.get("text", "")).strip()
            key = (normalize_text(heading).lower(), int(item.get("position", 0) or 0))
            if not heading or key in seen:
                continue
            seen.add(key)
            score, rejected, reason = score_objective_heading(heading, chapter)
            if score <= 0:
                continue
            candidates.append(
                {
                    "heading": heading,
                    "chapter": chapter,
                    "source_file": source_file,
                    "position": item.get("position", ""),
                    "page": item.get("page", ""),
                    "score": score,
                    "heading_score": score,
                    "heading_page": item.get("page", ""),
                    "rejected": rejected,
                    "heading_reject_reason": reason if rejected else "",
                    "rejected_reason": reason if rejected else "",
                }
            )
    candidates.sort(
        key=lambda item: (
            bool(item.get("rejected", False)),
            -int(item.get("score", 0) or 0),
            not is_chapter_one_label(str(item.get("chapter", ""))),
            int(item.get("position", 999999) or 999999),
        )
    )
    return candidates


def score_objective_heading(text: str) -> tuple[int, bool, str]:
    return score_objective_heading(text, "Unknown")


def score_objective_heading(text: str, chapter: str = "Unknown") -> tuple[int, bool, str]:
    cleaned = normalize_text(text).strip()
    lowered = cleaned.lower()
    if not cleaned:
        return 0, False, ""
    if len(cleaned) >= 120:
        return 10, True, "heading is too long to be a focused objective section heading"
    if cleaned.endswith("."):
        return 10, True, "heading ends with a full stop and appears narrative"
    if UNIVERSAL_FALSE_OBJECTIVE_PATTERN.search(cleaned):
        return 10, True, "non-research objective heading"
    for pattern in OBJECTIVE_STRONG_REJECT_PATTERNS:
        if pattern.search(cleaned):
            return 10, True, "RO marker, ordinal objective, findings, support, or discussion objective text"
    if not is_standalone_objective_heading_line(cleaned):
        return 10, True, "line appears narrative rather than a standalone objective heading"
    if any(term in lowered for term in OBJECTIVE_CONTEXT_REJECT_TERMS) and re.search(r"\b(objektif|objective|objectives)\b", lowered):
        return 10, True, "discussion, alignment, evaluation, or achievement heading"

    for phrase, priority in OBJECTIVE_HEADING_PRIORITIES:
        if phrase not in lowered:
            continue
        extra_words = count_heading_extra_words(lowered, phrase)
        if extra_words > 4:
            return 10, True, "objective phrase appears inside a long non-section heading"
        exact_bonus = 25 if lowered == phrase and phrase in EXACT_OBJECTIVE_HEADINGS else 0
        return priority + exact_bonus + objective_chapter_score_adjustment(chapter), False, ""

    if re.search(r"\b(objektif|objective|objectives)\b", lowered):
        return 10, True, "generic objective mention is not a ranked objective heading"
    return 0, False, ""


def is_standalone_objective_heading_line(text: str) -> bool:
    lowered = text.lower()
    if re.search(r"\b(menyatakan|menunjukkan|membuktikan|menghuraikan|menyokong|berdasarkan|dalam kalangan|dapatan)\b", lowered):
        return False
    if re.search(r"\b(this|these|that|which|yang)\b", lowered) and re.search(r"\b(objektif|objective|objectives)\b", lowered):
        return False
    if re.search(r"\([^)]*\bro\s*[0-9]+[^)]*\)", lowered, re.IGNORECASE):
        return False
    return True


def objective_chapter_score_adjustment(chapter: str) -> int:
    lowered = chapter.lower()
    if re.search(r"\b(bab|chapter)\s*(1|i)\b", lowered):
        return 100
    if re.search(r"\b(bab|chapter)\s*(2|ii)\b", lowered):
        return 20
    if re.search(r"\b(bab|chapter)\s*(3|iii)\b", lowered):
        return 10
    if re.search(r"\b(bab|chapter)\s*(4|iv|5|v)\b", lowered):
        return -50
    return 0


def count_heading_extra_words(heading: str, phrase: str) -> int:
    heading_words = re.findall(r"[a-zA-ZÀ-ž0-9]+", heading)
    phrase_words = re.findall(r"[a-zA-ZÀ-ž0-9]+", phrase)
    return max(0, len(heading_words) - len(phrase_words))


def infer_document_chapter(document: dict[str, Any]) -> str:
    source_file = str(document.get("source_file", ""))
    chapter_match = re.search(r"\b(bab|chapter)\s*([0-9ivx]+)\b", source_file, re.IGNORECASE)
    if chapter_match:
        label = chapter_match.group(1).title()
        return f"{label} {chapter_match.group(2).upper()}"
    chapters = [item for item in document.get("chapters", []) if isinstance(item, dict)]
    if chapters:
        label = str(chapters[0].get("label", "")).strip()
        if label:
            return label
    return "Unknown"


def is_chapter_one_label(label: str) -> bool:
    return bool(re.search(r"\b(bab|chapter)\s*(1|i)\b", label, re.IGNORECASE))


def collect_universal_objective_section(
    source_file: str,
    selected_heading: dict[str, Any],
    paragraphs: list[dict[str, Any]],
    headings: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    heading_position = int(selected_heading.get("position", 0) or 0)
    stop_position = find_universal_stop_position(source_file, heading_position, headings)
    section = []
    for paragraph in paragraphs:
        if str(paragraph.get("source_file", source_file)) != source_file:
            continue
        position = int(paragraph.get("position", 0) or 0)
        if position <= heading_position:
            continue
        if stop_position and position >= stop_position:
            break
        text = str(paragraph.get("text", ""))
        if UNIVERSAL_STOP_HEADING_PATTERN.search(text):
            break
        section.append({**paragraph, "text": cleanup_pdf_fragment(text)})
    return [item for item in section if item.get("text")]


def find_universal_stop_position(source_file: str, heading_position: int, headings: list[dict[str, Any]]) -> int | None:
    stops = []
    for heading in headings:
        if str(heading.get("source_file", source_file)) != source_file:
            continue
        position = int(heading.get("position", 0) or 0)
        text = str(heading.get("text", ""))
        if position <= heading_position:
            continue
        if UNIVERSAL_STOP_HEADING_PATTERN.search(text) or is_likely_major_heading(text):
            stops.append(position)
    return min(stops) if stops else None


def extract_general_objective(
    source_file: str,
    selected_heading: dict[str, Any],
    section: list[dict[str, Any]],
) -> dict[str, Any]:
    for index, item in enumerate(section):
        text = str(item.get("text", ""))
        if not UNIVERSAL_GENERAL_SUBHEADING_PATTERN.match(text):
            continue
        next_items = section[index + 1 :]
        for candidate in next_items:
            candidate_text = str(candidate.get("text", ""))
            if UNIVERSAL_SPECIFIC_SUBHEADING_PATTERN.match(candidate_text):
                return {}
            cleaned = strip_objective_prefix(candidate_text)
            if is_objective_sentence(cleaned, allow_without_prefix=True):
                return {
                    "objective_text": cleaned,
                    "source_file": source_file,
                    "source_chapter": "Bab 1",
                    "source_heading": str(selected_heading.get("text", "Research Objectives")),
                    "confidence_score": 60,
                    "extraction_status": "general_objective",
                    "objective_source_page": candidate.get("page", ""),
                }
    return {}


def choose_specific_section(section: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], dict[str, Any] | None]:
    for index, item in enumerate(section):
        if UNIVERSAL_SPECIFIC_SUBHEADING_PATTERN.match(str(item.get("text", ""))):
            return section[index + 1 :], item
    return section, None


def extract_specific_objectives(
    source_file: str,
    source_heading: str,
    selected_subheading: dict[str, Any] | None,
    section: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], str, str, list[dict[str, Any]]]:
    rejected: list[dict[str, Any]] = []
    phase_items, phase_rejected = extract_objectives_by_pattern(source_file, source_heading, selected_subheading, section, "phase")
    rejected.extend(phase_rejected)
    if phase_items:
        return phase_items, "phase_objectives", "phase", rejected

    numbered_items, numbered_rejected = extract_objectives_by_pattern(source_file, source_heading, selected_subheading, section, "numbered")
    rejected.extend(numbered_rejected)
    if numbered_items:
        return numbered_items, "numbered_objectives", "numbered", rejected

    bullet_items, bullet_rejected = extract_objectives_by_pattern(source_file, source_heading, selected_subheading, section, "bullet")
    rejected.extend(bullet_rejected)
    if bullet_items:
        return bullet_items, "bullet_objectives", "bullet", rejected

    return [], "none", "", rejected


def extract_objectives_by_pattern(
    source_file: str,
    source_heading: str,
    selected_subheading: dict[str, Any] | None,
    section: list[dict[str, Any]],
    mode: str,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    items = []
    rejected = []
    for paragraph in section:
        text = str(paragraph.get("text", ""))
        if is_subheading_line(text):
            continue
        item_text = ""
        confidence = 60
        if mode == "phase":
            match = UNIVERSAL_PHASE_PATTERN.match(text)
            if match:
                item_text = strip_objective_prefix(match.group(2))
                confidence = 90
        elif mode == "numbered":
            match = UNIVERSAL_NUMBERED_PATTERN.match(text)
            if match:
                item_text = strip_objective_prefix(match.group(5))
                confidence = 95
        elif mode == "bullet":
            match = UNIVERSAL_BULLET_PATTERN.match(text)
            if match:
                item_text = strip_objective_prefix(match.group(1))
                confidence = 80

        if not item_text:
            if looks_like_possible_objective_noise(text):
                rejected.append(reject_candidate(paragraph, "narrative or unnumbered objective-like paragraph"))
            continue
        valid, reason = validate_universal_objective(item_text, mode)
        if not valid:
            rejected.append(reject_candidate(paragraph, reason))
            continue
        items.append(
            {
                "objective_text": item_text,
                "source_file": source_file,
                "source_chapter": "Bab 1",
                "source_heading": source_heading,
                "confidence_score": confidence,
                "extraction_status": "extracted",
                "selected_heading": source_heading,
                "selected_subheading": str(selected_subheading.get("text", "")) if selected_subheading else "",
                "objective_source_page": paragraph.get("page", ""),
                "position": int(paragraph.get("position", 0) or 0),
            }
        )
    return items, rejected


def validate_universal_objective(text: str, mode: str) -> tuple[bool, str]:
    if not text:
        return False, "empty objective text"
    if UNIVERSAL_FALSE_OBJECTIVE_PATTERN.search(text):
        return False, "non-research objective"
    if looks_like_citation_fragment(text):
        return False, "citation fragment"
    words = text.split()
    if len(words) < 8 and mode != "phase":
        return False, "incomplete fragment shorter than 8 words"
    if ends_with_broken_author_citation(text):
        return False, "text ends with broken author citation"
    if not UNIVERSAL_OBJECTIVE_VERB_PATTERN.search(text):
        return False, "research objective verb not detected"
    return True, ""


def cleanup_pdf_fragment(text: str) -> str:
    cleaned = normalize_text(text)
    cleaned = re.sub(r"^\(?[A-Z][A-Za-zÀ-ÿ'`\-]+,\s*(?:19|20)\d{2}[a-z]?\)?\.?$", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def strip_objective_prefix(text: str) -> str:
    stripped = normalize_text(text)
    stripped = re.sub(r"^(?:untuk|to)\s+", "", stripped, flags=re.IGNORECASE)
    return stripped.strip(" ;:-")


def is_objective_sentence(text: str, allow_without_prefix: bool = False) -> bool:
    if not text or UNIVERSAL_FALSE_OBJECTIVE_PATTERN.search(text):
        return False
    if len(text.split()) < 8:
        return False
    return allow_without_prefix or bool(UNIVERSAL_OBJECTIVE_VERB_PATTERN.search(text))


def is_subheading_line(text: str) -> bool:
    return bool(UNIVERSAL_GENERAL_SUBHEADING_PATTERN.match(text) or UNIVERSAL_SPECIFIC_SUBHEADING_PATTERN.match(text))


def looks_like_possible_objective_noise(text: str) -> bool:
    return bool(re.search(r"\b(?:objective|objektif|tujuan|matlamat|identify|mengenal)\b", text, re.IGNORECASE))


def ends_with_broken_author_citation(text: str) -> bool:
    return bool(re.search(r"\([A-Z][A-Za-zÀ-ÿ'`\-]+(?:\s+et\s+al\.)?,?\s*$", text))


def reject_candidate(candidate: dict[str, Any], reason: str) -> dict[str, Any]:
    return {
        "source_file": str(candidate.get("source_file", "Unknown file")),
        "text": str(candidate.get("text", "")),
        "position": candidate.get("position", 0),
        "page": candidate.get("page", ""),
        "rejected_reason": reason,
    }


def reject_objective_candidates(candidates: list[dict[str, Any]], reason: str) -> list[dict[str, Any]]:
    return [reject_candidate(candidate, reason) for candidate in candidates if isinstance(candidate, dict)]
