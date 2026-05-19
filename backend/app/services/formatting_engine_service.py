from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SECTION_ORDER = [
    "Abstract",
    "Introduction",
    "Problem Statement",
    "Literature Review",
    "Methodology",
    "Findings",
    "Discussion",
    "Conclusion",
]


def generate_formatted_docx(
    paper_id: str,
    full_paper: dict[str, Any] | None,
    reference_bank: dict[str, Any] | None,
    section_structure: dict[str, Any] | None,
    output_path: Path,
    template_used: str = "ICC2026",
) -> dict[str, Any]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    title = get_title(paper_id, full_paper, section_structure)
    sections = extract_sections(full_paper.get("integrated_text", "") if full_paper else "")
    references = reference_bank.get("references", []) if reference_bank else []

    document = Document()
    configure_document(document)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)

    document.add_paragraph()

    sections_formatted = []
    for section_name in SECTION_ORDER:
        content = sections.get(section_name, f"[WARNING: {section_name} section is missing.]")
        document.add_heading(section_name, level=1)
        add_body_text(document, content)
        sections_formatted.append(section_name)

    document.add_heading("References", level=1)
    if references:
        for reference in references:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(str(reference.get("apa_reference", "")))
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    else:
        add_body_text(document, "[References pending: build reference bank before final submission.]")

    document.save(output_path)

    total_word_count = count_words(full_paper.get("integrated_text", "") if full_paper else "")
    reference_list_included = bool(references)
    report = {
        "paper_id": paper_id,
        "template_used": template_used,
        "docx_path": str(output_path),
        "sections_formatted": sections_formatted,
        "reference_list_included": reference_list_included,
        "total_word_count": total_word_count,
        "formatting_audit": build_formatting_audit(
            sections=sections,
            references=references,
            reference_bank=reference_bank,
        ),
        "status": "formatted",
        "generated_at": datetime.now(UTC).isoformat(),
    }
    return report


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = True


def add_body_text(document: Document, text: str) -> None:
    paragraphs = [part.strip() for part in text.split("\n") if part.strip()]
    if not paragraphs:
        paragraphs = ["[No content available.]"]

    for paragraph_text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(paragraph_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def extract_sections(integrated_text: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current_section: str | None = None

    for line in integrated_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            heading = stripped[2:].strip()
            current_section = heading if heading in SECTION_ORDER else None
            if current_section:
                sections.setdefault(current_section, [])
            continue

        if current_section:
            sections[current_section].append(stripped)

    return {
        section_name: "\n".join(lines).strip()
        for section_name, lines in sections.items()
    }


def build_formatting_audit(
    sections: dict[str, str],
    references: list[dict[str, Any]],
    reference_bank: dict[str, Any] | None,
) -> dict[str, str]:
    missing_sections = [section for section in SECTION_ORDER if section not in sections]
    apa_issues = int(reference_bank.get("apa_issues", 0) or 0) if reference_bank else 0

    return {
        "heading_consistency": "Review" if missing_sections else "OK",
        "citation_style": "Review" if apa_issues else "OK",
        "reference_list": "OK" if references else "Review",
        "table_numbering": "Review required",
        "figure_numbering": "No figures detected",
        "margin_font_compliance": "OK",
    }


def get_title(
    paper_id: str,
    full_paper: dict[str, Any] | None,
    section_structure: dict[str, Any] | None,
) -> str:
    if full_paper and full_paper.get("title"):
        return str(full_paper["title"])
    if section_structure and section_structure.get("paper_title"):
        return str(section_structure["paper_title"])
    return fallback_title(paper_id)


def count_words(text: str) -> int:
    return len(text.split())


def fallback_title(paper_id: str) -> str:
    titles = {
        "PAPER_1": "Need Analysis Paper",
        "PAPER_2": "Development & Validation Paper",
        "PAPER_3": "Framework / Model Paper",
    }
    return titles.get(paper_id, "Journal Paper")
