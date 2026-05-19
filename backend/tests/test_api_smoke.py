from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.api import audit, citation, intelligence, journal, knowledge_graph, objective, parser, table, upload, workflow
from app.api.upload import write_upload_metadata
from app.main import app


PROJECT_ID = "PROJECT_001"


@pytest.fixture()
def client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> TestClient:
    upload_root = tmp_path / "storage" / "uploads"
    generated_root = tmp_path / "storage" / "generated_outputs"
    formatted_root = tmp_path / "storage" / "formatted_outputs"

    monkeypatch.setattr(upload, "UPLOAD_ROOT", upload_root)
    monkeypatch.setattr(parser, "GENERATED_OUTPUT_ROOT", generated_root)
    monkeypatch.setattr(intelligence, "GENERATED_OUTPUT_ROOT", generated_root)
    monkeypatch.setattr(citation, "GENERATED_OUTPUT_ROOT", generated_root)
    monkeypatch.setattr(objective, "GENERATED_OUTPUT_ROOT", generated_root)
    monkeypatch.setattr(table, "GENERATED_OUTPUT_ROOT", generated_root)
    monkeypatch.setattr(audit, "GENERATED_OUTPUT_ROOT", generated_root)
    monkeypatch.setattr(knowledge_graph, "GENERATED_OUTPUT_ROOT", generated_root)
    monkeypatch.setattr(journal, "GENERATED_OUTPUT_ROOT", generated_root)
    monkeypatch.setattr(journal, "FORMATTED_OUTPUT_ROOT", formatted_root)
    monkeypatch.setattr(workflow, "GENERATED_OUTPUT_ROOT", generated_root)

    create_project_uploads(upload_root / PROJECT_ID)
    return TestClient(app)


def create_project_uploads(project_dir: Path) -> None:
    project_dir.mkdir(parents=True, exist_ok=True)
    metadata = {}

    for chapter in ["Bab 1", "Bab 2", "Bab 3", "Bab 4", "Bab 5"]:
        document = Document()
        document.add_heading(chapter, level=1)
        document.add_paragraph(
            f"Objektif kajian {chapter}: membangunkan modul dakwah berstruktur (Ali, 2020)."
        )
        if chapter == "Bab 3":
            document.add_heading("Methodology", level=2)
            document.add_paragraph("Kajian ini menggunakan pendekatan DDR dan pemetaan aktiviti.")
        if chapter == "Bab 4":
            document.add_heading("Findings", level=2)
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Aktiviti"
            table.cell(0, 1).text = "Kompetensi"
            table.cell(1, 0).text = "Pemilihan kandungan"
            table.cell(1, 1).text = "Pemetaan aktiviti"
            document.add_paragraph("Jadual 4.1 Pemetaan aktiviti dakwah")
        if chapter == "Bab 5":
            document.add_heading("Discussion", level=2)
            document.add_heading("References", level=1)
            document.add_paragraph("Ali, A. (2020). Kajian dakwah akademik.")

        file_path = project_dir / f"{chapter.replace(' ', '_')}.docx"
        document.save(file_path)
        metadata[file_path.name] = {
            "file_type": "thesis_chapter",
            "chapter_label": chapter,
        }

    mfl_path = project_dir / "MFL.xlsx"
    mfl_path.write_bytes(b"placeholder")
    metadata[mfl_path.name] = {"file_type": "mfl", "chapter_label": "MFL"}
    write_upload_metadata(PROJECT_ID, metadata)


def assert_wrapped_success(response) -> dict:
    assert response.status_code == 200
    payload = response.json()
    assert payload["success"] is True
    assert payload["status"] == "success"
    assert "data" in payload
    return payload["data"]


def test_health_endpoint_is_wrapped(client: TestClient) -> None:
    data = assert_wrapped_success(client.get("/health"))
    assert data["status"] == "ok"
    assert data["service"] == "Thesis2Journal AI API"


def test_projects_endpoints_are_wrapped(client: TestClient) -> None:
    projects_data = assert_wrapped_success(client.get("/projects"))
    assert projects_data["projects"]

    project_data = assert_wrapped_success(client.get(f"/projects/{PROJECT_ID}"))
    assert project_data["project"]["id"] == PROJECT_ID


def test_full_pipeline_and_status_are_wrapped(client: TestClient) -> None:
    pipeline_data = assert_wrapped_success(client.post(f"/workflow/{PROJECT_ID}/run-full-pipeline"))
    assert pipeline_data["pipeline_status"] == "completed"
    assert pipeline_data["completed_count"] >= 30
    assert pipeline_data["final_download_urls"]["docx"]
    assert pipeline_data["final_download_urls"]["markdown"]

    status_data = assert_wrapped_success(client.get(f"/workflow/{PROJECT_ID}/status"))
    assert status_data["pipeline_status"] == "completed"
    assert status_data["generated_files_summary"]


def test_missing_workflow_status_returns_wrapped_404(client: TestClient) -> None:
    response = client.get("/workflow/MISSING_PROJECT/status")
    assert response.status_code == 404
    payload = response.json()
    assert payload["success"] is False
    assert payload["status"] == "error"
    assert "data" in payload
